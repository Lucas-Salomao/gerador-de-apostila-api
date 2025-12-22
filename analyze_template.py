"""
Script para analisar o template DOCX e identificar placeholders.
"""
from docx import Document
from pathlib import Path

template_path = Path('template/Template Docx.docx')
doc = Document(template_path)

print('=' * 60)
print('ANÁLISE DO TEMPLATE DOCX')
print('=' * 60)

print('\n📌 PLACEHOLDERS ENCONTRADOS:')
placeholders = []
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if '{{' in text and '}}' in text:
        placeholders.append({
            'index': i,
            'style': para.style.name,
            'text': text
        })
        print(f'  [{i}] Estilo: "{para.style.name}"')
        print(f'       Texto: "{text}"')
        print()

if not placeholders:
    print('  Nenhum placeholder encontrado!')

print('\n📄 ESTRUTURA DO DOCUMENTO:')
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if text:
        display_text = text[:60] + '...' if len(text) > 60 else text
        style = para.style.name
        
        # Destacar placeholders
        if '{{' in text:
            print(f'  🔹 [{i}] {style}: {display_text}')
        elif 'Heading' in style or 'Title' in style:
            print(f'  📗 [{i}] {style}: {display_text}')
        else:
            print(f'     [{i}] {style}: {display_text}')

print('\n📊 RESUMO:')
print(f'  Total de parágrafos: {len(doc.paragraphs)}')
print(f'  Total de seções: {len(doc.sections)}')
print(f'  Placeholders: {len(placeholders)}')
for p in placeholders:
    print(f'    - {p["text"]}')
