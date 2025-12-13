import os
import tempfile
from typing import Dict, List, Tuple, Any, TypedDict, Optional
import json
from pathlib import Path
import logging
import re
from dotenv import load_dotenv
import time

# Bibliotecas para LangGraph
from langgraph.graph import StateGraph, END
from langgraph.checkpoint.memory import MemorySaver
import langchain

# Bibliotecas para Gemini/Vertex AI
# Import condicional para suportar ambos os backends
USE_VERTEXAI = os.getenv("USE_VERTEXAI", "false").lower() == "true"

if USE_VERTEXAI:
    from google.cloud import aiplatform
    from vertexai.generative_models import GenerativeModel
    import vertexai
else:
    import google.generativeai as genai

# Biblioteca para exportação
import docx
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

load_dotenv()

# # Configuração de logs
logging.basicConfig(
    level=logging.DEBUG,
    format="%(asctime)s [%(levelname)s] %(message)s",
    handlers=[
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

# Inicializar Gemini API
def init_gemini_api():
    """Inicializa a conexão com o Gemini API."""
    logger.info("Inicializando Gemini API...")
    import google.generativeai as genai
    genai.configure(api_key=os.getenv("GEMINI_API_KEY"))
    return genai.GenerativeModel(os.getenv("GEMINI_MODEL"))

def init_vertex_ai():
    """Inicializa a conexão com o Vertex AI usando credenciais do ambiente."""
    logger.info("Inicializando Vertex AI...")
    from google.cloud import aiplatform
    from vertexai.generative_models import GenerativeModel
    import vertexai
    
    project_id = os.getenv("GOOGLE_CLOUD_PROJECT")
    location = os.getenv("GOOGLE_CLOUD_LOCATION", "us-central1")
    
    if not project_id:
        raise ValueError("A variável de ambiente GOOGLE_CLOUD_PROJECT não foi definida.")
    
    vertexai.init(project=project_id, location=location)
    
    # Modelo no Vertex AI
    return GenerativeModel(os.getenv("GEMINI_MODEL"))

def get_model():
    """Retorna o modelo correto baseado na variável USE_VERTEXAI."""
    if USE_VERTEXAI:
        logger.info("Usando Vertex AI como backend")
        return init_vertex_ai()
    else:
        logger.info("Usando Gemini API como backend")
        return init_gemini_api()

# Função auxiliar para parsing seguro de JSON
def safe_json_parse(response_text: str, fallback: Any) -> Any:
    """Tenta decodificar JSON e retorna um fallback em caso de erro."""
    if response_text.startswith("```json"):
        response_text = response_text[7:]
        if response_text.endswith("```"):
            response_text = response_text[:-3]
    elif response_text.startswith("```"):
        response_text = response_text[3:]
        if response_text.endswith("```"):
            response_text = response_text[:-3]
    
    try:
        return json.loads(response_text)
    except json.JSONDecodeError:
        logger.error(f"Erro ao decodificar JSON: {response_text[:100]}... Usando fallback.")
        return fallback

# Definição dos estados do grafo
class BookState(TypedDict, total=False):
    theme: str
    title: str
    area_tecnologica: str
    target_audience: str
    num_chapters: int
    outline: List[Dict[str, Any]]
    chapters: Dict[int, Dict[str, str]]
    current_chapter: int
    status: str
    feedback: str
    export_path: str
    feedback_path: str

# Funções para cada etapa do processo
def get_book_info(state: BookState, model) -> Dict[str, Any]:
    """Obtém informações básicas e gera o título com base no tema."""
    logger.info(f"Estado recebido em get_book_info: {state}")
    logger.info("Coletando informações básicas do livro...")
    updates = {}
    
    theme = state.get("theme", "Um tema genérico")
    area_tecnologica = state.get("area_tecnologica", "Não especificada")
    target_audience = state.get("target_audience", "Adultos")
    
    updates["theme"] = theme
    updates["area_tecnologica"] = area_tecnologica
    updates["target_audience"] = target_audience
    
    prompt = f"""
    Você é um especialista em redação técnica. Baseado no seguinte tema, área tecnológica e público-alvo, sugira um título formal e técnico que reflita um enfoque analítico e informativo:
    Tema: {theme}
    Área Tecnológica: {area_tecnologica}
    Público-Alvo: {target_audience}
    Responda SOMENTE em formato JSON com a chave "title", sem texto adicional. Exemplo: {{"title": "Fundamentos de Exploração Espacial"}}. Não inclua bloco de código, ou seja ```json```
    O Título deve ter no máximo 80 caracteres. Caracteres inválidos para o título: , \\ / : * ? " < > |
    """
    
    logger.info("Gerando título com base no tema...")
    response = generate_with_retry(model, prompt)
    logger.debug(f"Resposta bruta do modelo: {response.text}")
    info = safe_json_parse(response.text, {"title": f"Livro sobre {theme}"})
    updates["title"] = info.get("title", f"Livro sobre {theme}")
    logger.info(f"Título gerado: {updates['title']}")
    
    updates["status"] = "book_info_collected"
    logger.debug(f"Atualizações de get_book_info: {updates}")
    return updates

def create_outline(state: BookState, model) -> Dict[str, Any]:
    """Cria o sumário do livro baseado nas informações fornecidas."""
    logger.info("Criando sumário do livro...")
    
    prompt = f"""
    Você é um especialista técnico no tema {state['theme']}, elaborando um livro técnico para estudo de um determinado tema. 
    Baseado nas seguintes informações, crie um sumário detalhado e com foco em aspectos técnicos e práticos:
    
    Tema: {state['theme']}
    Título sugerido: {state['title']}
    Área Tecnológica: {state.get('area_tecnologica', 'Não especificada')}
    Público-Alvo: {state['target_audience']}
    
    Cada capítulo deve ter uma numeração inteira e sequencial (exemplo: 1, 2, 3, 4, 5, etc.)

    O sumário deve conter exatamente {state['num_chapters']} capítulos, cada um abordando um aspecto técnico ou prático do tema, com títulos objetivos e descrições que detalhem o conteúdo analítico a ser explorado.
    Responda SOMENTE em formato JSON com uma lista de objetos contendo "chapter_number", "chapter_title" e "chapter_description".
    Exemplo: [{{"chapter_number": 1, "chapter_title": "Princípios de Propulsão Espacial", "chapter_description": "Análise dos sistemas de propulsão usados em missões espaciais"}}]
    Não inclua bloco de código, ou seja ```json```
    """

    response = generate_with_retry(model, prompt)
    if not response:
        logger.error("Falha ao gerar sumário.")
        return {"status": "error", "message": "Falha ao gerar sumário."}

    logger.debug(f"Resposta bruta do modelo: {response.text}")
    outline_data = safe_json_parse(response.text, [
        {"chapter_number": 1, "chapter_title": "Introdução", 
         "chapter_description": f"Exploração inicial do tema {state['theme']}."}
    ])
    
    num_chapters = state.get("num_chapters", 5)
    if len(outline_data) < num_chapters:
        logger.warning(f"Sumário com menos de {num_chapters} capítulos. Adicionando capítulos extras.")
        for i in range(len(outline_data) + 1, num_chapters + 1):
            outline_data.append({
                "chapter_number": i,
                "chapter_title": f"Capítulo {i}",
                "chapter_description": f"Continuação da exploração de {state['theme']}."
            })
    
    updates = {
        "outline": outline_data,
        "chapters": {item["chapter_number"]: {"title": item["chapter_title"], 
                                              "description": item["chapter_description"],
                                              "content": ""} 
                     for item in outline_data},
        "current_chapter": 1,
        "status": "outline_created"
    }
    logger.info(f"Sumário criado com {len(outline_data)} capítulos.")
    logger.info(f"Capítulos gerados: {updates['chapters']}")
    return updates

def write_chapter(state: BookState, model, st_session=None) -> Dict[str, Any]:
    """Escreve o conteúdo para o capítulo atual."""
    current = state["current_chapter"]
    updates = {}
    if current > len(state["chapters"]):
        updates["status"] = "all_chapters_written"
        logger.info("Todos os capítulos foram escritos.")
        return updates
    
    chapter_info = state["chapters"][current]
    logger.info(f"Escrevendo Capítulo {current}: {chapter_info['title']}...")
        
    prev_content = ""
    if current > 1 and state["chapters"].get(current-1, {}).get("content"):
        prev_chapter = state["chapters"][current-1]
        prev_content = f"""
        Resumo do capítulo anterior ({current-1}: {prev_chapter['title']}):
        {prev_chapter['content'][:500]}... (resumido)
        """
    
    prompt = f"""
    Você é um especialista técnico escrevendo um livro intitulado "{state['title']}" com o tema "{state['theme']}".
    A área tecnológica do livro é "{state.get('area_tecnologica', 'Não especificada')}", direcionado para o público "{state['target_audience']}"
    
    Escreva o Capítulo {current}: "{chapter_info['title']}".
    
    Descrição do capítulo: {chapter_info['description']}
    
    {prev_content}
    
    Escreva um texto técnico e analítico, com linguagem formal e objetiva. Inclua informações técnicas detalhadas, exemplos contextualizados (reais ou hipotéticos), dados relevantes e explicações claras. Evite diálogos narrativos ou descrições literárias excessivas. Estruture o conteúdo com seções claras (ex.: introdução, desenvolvimento, análise, exemplos, conclusão). O capítulo deve ter pelo menos 3000 palavras. Seja o mais detalhista e técnico possível e aborde o tema do capítulo com profundidade técnica e bastante exemplos.
    Estruture o capítulo com títulos e subtítulos para facilitar a leitura e compreensão do conteúdo. Siga a numeração do capítulo e estruture os subtítulos com base na numeração do capítulo.
    """

    response = generate_with_retry(model, prompt)

    if not response:
        logger.error("Falha ao gerar conteúdo para o capítulo.")
        return {"status": "error", "message": "Falha ao gerar conteúdo para o capítulo."}

    updated_chapters = state["chapters"].copy()
    updated_chapters[current]["content"] = response.text
    updates["chapters"] = updated_chapters
    logger.info(f"Capítulo {current} concluído com sucesso.")
    
    # Exibir o conteúdo gerado no Streamlit
    if st_session:
        # st_session.write(f"### Capítulo {current}: {chapter_info['title']}")
        st_session.write(response.text)
    time.sleep(5)
    updates["current_chapter"] = current + 1
    updates["status"] = "chapter_written" if updates["current_chapter"] <= len(state["chapters"]) else "all_chapters_written"
    return updates

def review_and_edit(state: BookState, model) -> Dict[str, Any]:
    """Revisa e edita o livro completo."""
    logger.info("Revisando e editando o livro...")
    book_summary = f"""
    Tema: {state['theme']}
    Título: {state['title']}
    Área Tecnológica: {state.get('area_tecnologica', 'Não especificada')}
    Público-alvo: {state['target_audience']}
    
    Sumário:
    """
    for chapter_num, chapter_data in sorted(state["chapters"].items()):
        book_summary += f"\nCapítulo {chapter_num}: {chapter_data['title']} - {chapter_data['description'][:100]}..."
    
    prompt = f"""
    Você é um editor revisando o livro:
    
    {book_summary}
    
    Forneça feedback sobre estrutura, fluxo narrativo, consistência com o tema "{state['theme']}" 
    e apelo ao público-alvo. Seja minucioso referente às informações técnicas e sugira melhorias. Revise tecnicamente o livro e verifique se há alguma inconsistência.
    Traga sugestões de melhorias, correções e ajustes necessários, indicando os capítulos e seções específicas para correção.
    """

    response = generate_with_retry(model, prompt)

    if not response:
        logger.error("Falha ao gerar feedback.")
        return {"status": "error", "message": "Falha ao gerar feedback."}

    updates = {
        "feedback": response.text,
        "status": "reviewed"
    }
    logger.info("Revisão concluída. Feedback gerado.")
    
    return updates

def export_feedback(state: BookState) -> Dict[str, Any]:
    """Processa o feedback (sem criar arquivo local - feedback já está embutido no DOCX)."""
    logger.info("Processando feedback...")
    # Não salvamos mais arquivo TXT local, o feedback é embutido no DOCX
    updates = {
        "feedback_path": None,  # Não há mais arquivo local
        "status": "feedback_exported"
    }
    logger.info("Feedback processado (embutido no documento final).")
    return updates

# def export_book(state: BookState) -> Dict[str, Any]:
#     """Exporta o livro apenas para DOCX."""
#     logger.info("Exportando livro para DOCX...")
#     doc = Document()
#     doc.add_heading(state["title"], 0)
#     doc.add_paragraph(f"Tema: {state['theme']}")
#     doc.add_paragraph(f"Gênero: {state['genre']}")
#     doc.add_paragraph(f"Público-alvo: {state['target_audience']}")
    
#     for chapter_num, chapter_data in sorted(state["chapters"].items()):
#         doc.add_heading(f"Capítulo {chapter_num}: {chapter_data['title']}", 1)
#         doc.add_paragraph(chapter_data["content"])
    
#     doc_path = f"{state['title'].replace(' ', '_')}.docx"
#     doc.save(doc_path)
#     updates = {
#         "export_path": doc_path,
#         "status": "exported"
#     }
#     logger.info(f"Livro exportado com sucesso para: {doc_path}")
#     return updates


def export_book(state: BookState) -> Dict[str, Any]:
    """Exporta o livro para DOCX com formatação completa e suporte total a Markdown."""
    logger.info("Exportando livro para DOCX com formatação completa...")
    doc = Document()

    # Configuração de estilos
    styles = doc.styles
    for style in styles:
        if hasattr(style, 'font') and style.font:
            style.font.name = 'Arial'

    title_style = styles['Title']
    title_style.font.name = 'Arial'
    title_style.font.size = Pt(24)
    title_style.font.bold = True
    title_style.font.color.rgb = RGBColor(0, 0, 0)

    heading_styles = {
        1: styles['Heading 1'], 2: styles['Heading 2'], 3: styles['Heading 3'],
        4: styles['Heading 4'], 5: styles['Heading 5'], 6: styles['Heading 6']
    }
    for level, style in heading_styles.items():
        style.font.name = 'Arial'
        style.font.size = Pt(16 - level * 2)  # Diminui o tamanho conforme o nível
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 51, 102)

    normal_style = styles['Normal']
    normal_style.font.name = 'Arial'
    normal_style.font.size = Pt(12)
    normal_style.paragraph_format.line_spacing = 1.15
    normal_style.paragraph_format.space_after = Pt(10)

    # Configurar margens e rodapé
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

        footer = section.footer
        footer_paragraph = footer.paragraphs[0]
        footer_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        footer_run = footer_paragraph.add_run()
        footer_run.font.name = 'Arial'
        footer_run.font.size = Pt(10)

        def add_field(run, field_code):
            field = OxmlElement('w:fldChar')
            field.set(qn('w:fldCharType'), 'begin')
            run._element.append(field)
            instr = OxmlElement('w:instrText')
            instr.set(qn('xml:space'), 'preserve')
            instr.text = field_code
            run._element.append(instr)
            sep = OxmlElement('w:fldChar')
            sep.set(qn('w:fldCharType'), 'separate')
            run._element.append(sep)
            end = OxmlElement('w:fldChar')
            end.set(qn('w:fldCharType'), 'end')
            run._element.append(end)

        footer_run.add_text('Página ')
        add_field(footer_run, 'PAGE')
        footer_run.add_text(' de ')
        add_field(footer_run, 'NUMPAGES')

    # Adicionar título
    title_paragraph = doc.add_paragraph(state["title"], style='Title')
    title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Informações iniciais
    info_paragraph = doc.add_paragraph()
    info_paragraph.add_run("Tema: ").bold = True
    info_paragraph.add_run(state["theme"])
    info_paragraph.add_run("\nÁrea Tecnológica: ").bold = True
    info_paragraph.add_run(state["area_tecnologica"])
    info_paragraph.add_run("\nPúblico-alvo: ").bold = True
    info_paragraph.add_run(state["target_audience"])
    doc.add_page_break()

    # Adicionar sumário
    doc.add_heading("Sumário", level=1)
    toc_paragraph = doc.add_paragraph()
    toc_run = toc_paragraph.add_run()
    toc_field = OxmlElement('w:fldChar')
    toc_field.set(qn('w:fldCharType'), 'begin')
    toc_run._element.append(toc_field)
    toc_instr = OxmlElement('w:instrText')
    toc_instr.set(qn('xml:space'), 'preserve')
    toc_instr.text = 'TOC \\o "1-3" \\h \\z \\u'  # Níveis 1 a 3
    toc_run._element.append(toc_instr)
    toc_sep = OxmlElement('w:fldChar')
    toc_sep.set(qn('w:fldCharType'), 'separate')
    toc_run._element.append(toc_sep)
    toc_end = OxmlElement('w:fldChar')
    toc_end.set(qn('w:fldCharType'), 'end')
    toc_run._element.append(toc_end)
    doc.add_paragraph("Insira o sumário automático manualamente.", style='Caption')
    doc.add_page_break()

    # Parser Markdown completo
    def process_markdown(text, doc):
        lines = text.split('\n')
        in_code_block = False
        code_lines = []
        in_table = False
        table_rows = []

        for i, line in enumerate(lines):
            line = line.rstrip()
            if not line and not in_code_block and not in_table:
                continue

            # Blocos de código
            if line.strip() == '```':
                if not in_code_block:
                    in_code_block = True
                    code_lines = []
                else:
                    in_code_block = False
                    code_paragraph = doc.add_paragraph()
                    code_paragraph.paragraph_format.left_indent = Inches(0.5)
                    code_paragraph.paragraph_format.right_indent = Inches(0.5)
                    run = code_paragraph.add_run('\n'.join(code_lines))
                    run.font.name = 'Courier New'
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(50, 50, 50)
                continue
            if in_code_block:
                code_lines.append(line)
                continue

            # Títulos
            if line.startswith('#'):
                level = min(line.count('#', 0, 6), 6)
                title_text = line.lstrip('#').strip()
                doc.add_heading(title_text, level=min(level, 6))
                continue

            # Linha horizontal
            if re.match(r'^\s*[-*_]{3,}\s*$', line):
                doc.add_paragraph().add_run().add_break(docx.enum.text.WD_BREAK.LINE)
                continue

            # Citações
            if line.startswith('>'):
                quote_paragraph = doc.add_paragraph()
                quote_paragraph.paragraph_format.left_indent = Inches(0.5)
                apply_inline_formatting(line.lstrip('>').strip(), quote_paragraph)
                continue

            # Listas ordenadas e não ordenadas
            if re.match(r'^\s*(\d+\.|-|\*|\+)\s+', line):
                indent_level = len(re.match(r'^\s*', line).group()) // 2
                list_item = re.sub(r'^\s*(\d+\.|-|\*|\+)\s+', '', line).strip()
                style = 'List Number' if re.match(r'^\s*\d+\.\s+', line) else 'List Bullet'
                paragraph = doc.add_paragraph(style=style)
                paragraph.paragraph_format.left_indent = Inches(0.25 * (indent_level + 1))
                apply_inline_formatting(list_item, paragraph)
                continue

            # Tabelas
            if line.startswith('|') and '|' in line[1:]:
                if not in_table:
                    in_table = True
                    table_rows = []
                row = [cell.strip() for cell in line.split('|')[1:-1]]
                if row and not re.match(r'^\s*-+\s*$', row[0]):  # Ignora linha de separação
                    table_rows.append(row)
                elif table_rows:  # Fim da tabela após separador
                    in_table = False
                    # ADICIONE ESTA VERIFICAÇÃO
                    if len(table_rows) > 0 and len(table_rows[0]) > 0:
                        table = doc.add_table(rows=len(table_rows), cols=len(table_rows[0]))
                        table.style = 'Table Grid'
                        for r_idx, row_data in enumerate(table_rows):
                            # Garante que todas as linhas tenham o mesmo número de colunas
                            if len(row_data) == len(table_rows[0]):
                                for c_idx, cell_text in enumerate(row_data):
                                    try:
                                        cell = table.rows[r_idx].cells[c_idx]
                                        apply_inline_formatting(cell_text, cell.paragraphs[0])
                                    except IndexError:
                                        logger.warning(f"Ignorando célula fora do alcance na tabela. Linha {r_idx}, Coluna {c_idx}")
                    else:
                        logger.warning("Ignorando tentativa de criar uma tabela vazia ou malformada.")
                    table_rows = [] # Limpa para a próxima tabela
                continue

            # Parágrafo simples
            paragraph = doc.add_paragraph()
            apply_inline_formatting(line, paragraph)

    # Função para aplicar formatação inline
    def apply_inline_formatting(text, paragraph):
        # Regex para capturar negrito, itálico, tachado e links
        pattern = r'(\*\*\*.*?\*\*\*|\*\*.*?\*\*|\*.*?\*|~~.*?~~|$$ .*? $$$$ .*? $$|_.*?_)'
        parts = re.split(pattern, text)
        
        for part in parts:
            if not part:
                continue
            run = paragraph.add_run()
            if part.startswith('***') and part.endswith('***'):
                run.text = part[3:-3]
                run.bold = True
                run.italic = True
            elif part.startswith('**') and part.endswith('**'):
                run.text = part[2:-2]
                run.bold = True
            elif part.startswith('*') and part.endswith('*'):
                run.text = part[1:-1]
                run.italic = True
            elif part.startswith('_') and part.endswith('_'):
                run.text = part[1:-1]
                run.italic = True
            elif part.startswith('~~') and part.endswith('~~'):
                run.text = part[2:-2]
                run.font.strike = True
            elif part.startswith('[') and ']' in part and '(' in part and part.endswith(')'):
                link_text = part[1:part.index(']')]
                link_url = part[part.index('(')+1:-1]
                run.text = link_text
                run.font.underline = True
                run.font.color.rgb = RGBColor(0, 0, 255)
                # python-docx não suporta hiperlinks diretamente; URL é apenas visual
            else:
                run.text = part

    # Adicionar capítulos
    for chapter_num, chapter_data in sorted(state["chapters"].items()):
        doc.add_heading(f"Capítulo {chapter_num}: {chapter_data['title']}", level=1)
        process_markdown(chapter_data["content"], doc)
        if chapter_num < len(state["chapters"]):
            doc.add_page_break()

    if state.get("feedback"):
        doc.add_page_break()
        doc.add_heading("Feedback da Revisão", level=1)
        process_markdown(state["feedback"], doc)

    # Salvar em arquivo temporário (não na raiz do projeto)
    # O arquivo será enviado para o GCS e depois apagado automaticamente
    temp_dir = tempfile.gettempdir()
    safe_title = state['title'].replace(' ', '_').replace('/', '_').replace('\\', '_')
    doc_path = os.path.join(temp_dir, f"{safe_title}.docx")
    doc.save(doc_path)

    updates = {
        "export_path": doc_path,
        "status": "exported"
    }
    logger.info(f"Livro exportado com sucesso para: {doc_path}")
    return updates


def router(state: BookState) -> str:
    """Decide o próximo estado."""
    status_map = {
        "start": "get_book_info",
        "book_info_collected": "create_outline",
        "outline_created": "write_chapter",
        "chapter_written": "write_chapter",
        "all_chapters_written": "review_and_edit",
        "reviewed": "export_feedback",
        "feedback_exported": "export_book",
        "exported": END
    }
    next_state = status_map.get(state["status"], END)
    logger.debug(f"Transição de estado: {state['status']} -> {next_state}")
    return next_state

def create_book_agent(model, st_session=None):
    """Cria o agente de geração de livros."""
    logger.info("Criando agente de geração de livros...")
    workflow = StateGraph(BookState)
    
    workflow.add_node("get_book_info", lambda state: get_book_info(state, model))
    workflow.add_node("create_outline", lambda state: create_outline(state, model))
    workflow.add_node("write_chapter", lambda state: write_chapter(state, model, st_session))
    workflow.add_node("review_and_edit", lambda state: review_and_edit(state, model))
    workflow.add_node("export_feedback", export_feedback)
    workflow.add_node("export_book", export_book)
    
    workflow.set_entry_point("get_book_info")
    
    workflow.add_conditional_edges("get_book_info", router)
    workflow.add_conditional_edges("create_outline", router)
    workflow.add_conditional_edges("write_chapter", router)
    workflow.add_conditional_edges("review_and_edit", router)
    workflow.add_conditional_edges("export_feedback", router)
    workflow.add_conditional_edges("export_book", router)
    
    memory = MemorySaver()
    return workflow.compile(checkpointer=memory)

def generate_with_retry(model, prompt, retries=3, delay=5):
    for i in range(retries):
        try:
            response = model.generate_content(prompt)
            return response
        except Exception as e:
            logger.warning(f"Erro na chamada da API (tentativa {i+1}/{retries}): {e}")
            time.sleep(delay)
    logger.error("Falha ao gerar conteúdo após múltiplas tentativas.")
    return None # Ou lançar uma exceção

def agent_book_generator(area_tecnologica: str = "", custom_audience: str = "", custom_theme: str = "", custom_num_chapters: int = 5):
    """Executa o agente de geração de livros e emite atualizações de progresso."""
    logger.info("Iniciando processo de geração de livro...")
    try:
        model = get_model()
        book_agent = create_book_agent(model)

        initial_state = BookState(status="start")
        # ... (código de preenchimento do initial_state) ...
        if custom_theme: initial_state["theme"] = custom_theme
        if area_tecnologica: initial_state["area_tecnologica"] = area_tecnologica
        if custom_audience: initial_state["target_audience"] = custom_audience
        initial_state["num_chapters"] = custom_num_chapters

        config = {"configurable": {"thread_id": "1"}, "recursion_limit": 1000}

        # Define o número total de etapas para a barra de progresso
        total_steps = 2 + custom_num_chapters + 2  # Título, Sumário, N Capítulos, Revisão, Exportação

        for output in book_agent.stream(initial_state, config=config):
            node_name = list(output.keys())[0] if output else "unknown"
            node_output = output.get(node_name, {})
            stage = node_output.get("status", "desconhecido")

            # --- LÓGICA DA BARRA DE PROGRESSO MELHORADA ---
            progress_update = None
            current_step = 0

            if stage == "book_info_collected":
                current_step = 1
                progress_update = {
                    "type": "progress",
                    "text": f"Etapa {current_step}/{total_steps}: Gerando título...",
                    "value": int((current_step / total_steps) * 100)
                }
            elif stage == "outline_created":
                current_step = 2
                progress_update = {
                    "type": "progress",
                    "text": f"Etapa {current_step}/{total_steps}: Sumário criado!",
                    "value": int((current_step / total_steps) * 100)
                }
                # Emitir progresso do sumário
                yield progress_update
                
                # Anunciar que vai começar a escrever o primeiro capítulo
                progress_update = {
                    "type": "progress",
                    "text": f"Etapa 3/{total_steps}: Escrevendo capítulo 1/{custom_num_chapters}...",
                    "value": int((2.5 / total_steps) * 100)
                }
            elif node_name == "write_chapter" and "chapters" in node_output:
                # current_chapter indica o PRÓXIMO capítulo a ser escrito
                # Então o capítulo que acabou de ser concluído é current_chapter - 1
                written_chapter_num = node_output.get("current_chapter", 1) - 1
                next_chapter_num = node_output.get("current_chapter", 1)
                total_chapters = len(node_output["chapters"])
                
                if written_chapter_num > 0:
                    current_step = 2 + written_chapter_num
                    # Mensagem de conclusão do capítulo atual
                    progress_update = {
                        "type": "progress",
                        "text": f"Etapa {current_step}/{total_steps}: Capítulo {written_chapter_num}/{total_chapters} concluído!",
                        "value": int((current_step / total_steps) * 100)
                    }
                    # Emitir progresso de conclusão
                    yield progress_update
                    
                    # Se ainda há capítulos a escrever, anunciar o próximo
                    if next_chapter_num <= total_chapters:
                        progress_update = {
                            "type": "progress",
                            "text": f"Etapa {current_step + 1}/{total_steps}: Escrevendo capítulo {next_chapter_num}/{total_chapters}...",
                            "value": int(((current_step + 0.5) / total_steps) * 100)
                        }
                    else:
                        progress_update = None
            elif stage == "reviewed":
                current_step = 2 + custom_num_chapters + 1
                progress_update = {
                    "type": "progress",
                    "text": f"Etapa {current_step}/{total_steps}: Revisando o conteúdo...",
                    "value": int((current_step / total_steps) * 100)
                }
            elif stage == "exported":
                current_step = 2 + custom_num_chapters + 2
                progress_update = {
                    "type": "progress",
                    "text": f"Etapa {current_step}/{total_steps}: Gerando documento final...",
                    "value": int((current_step / total_steps) * 100)
                }
            
            if progress_update:
                yield progress_update
            # --- FIM DA LÓGICA DA BARRA DE PROGRESSO ---

            # O restante do código de 'yield' para o conteúdo de texto continua o mesmo
            # yield f"**Etapa Concluída:** {stage}"

            if stage == "book_info_collected":
                title = node_output.get('title', 'Título não gerado')
                yield f"# {title}\n"

            elif stage == "outline_created":
                outline = node_output.get('outline', [])
                if outline:
                    summary_text = "## Sumário\n---\n"
                    for item in outline:
                        num = item.get('chapter_number', 'N/A')
                        title = item.get('chapter_title', 'Sem título')
                        desc = item.get('chapter_description', 'Sem descrição')
                        summary_text += f"### Capítulo {num}: {title}\n"
                        summary_text += f"*{desc}*\n\n"
                    yield summary_text

            # Modificação para exibir o último capítulo
            elif stage == "chapter_written" or (node_name == "write_chapter" and stage == "all_chapters_written"):
                # O 'current_chapter' no estado já foi incrementado, então subtraímos 1 para pegar o que acabou de ser escrito.
                current_chap_num = node_output.get('current_chapter', 0) - 1
                if current_chap_num > 0:
                    chapters_data = node_output.get("chapters", {})
                    chapter_info = chapters_data.get(current_chap_num)

                    # Garante que a informação do capítulo e o conteúdo existam
                    if chapter_info and chapter_info.get("content"):
                        yield chapter_info["content"]

        checkpoint = book_agent.checkpointer.get(config)
        logger.info("Processo de geração de livro concluído!")
        
        # Extrair o estado correto do checkpoint
        # O checkpointer retorna um objeto com 'channel_values' contendo o estado
        if checkpoint and hasattr(checkpoint, 'channel_values'):
            final_state = dict(checkpoint.channel_values)
        elif isinstance(checkpoint, dict):
            final_state = checkpoint.get('channel_values', checkpoint)
        else:
            final_state = {}
        
        logger.info(f"DEBUG - final_state keys: {final_state.keys() if isinstance(final_state, dict) else 'not a dict'}")
        logger.info(f"DEBUG - export_path in final_state: {final_state.get('export_path', 'NOT FOUND')}")

        yield {"final_state": final_state}

    except Exception as e:
        logger.error(f"Erro durante a geração do livro: {e}")
        yield {"final_state": {"status": "error", "message": str(e)}}